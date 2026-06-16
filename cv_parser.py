"""
cv_parser.py
================
Extracts raw text from an uploaded CV (PDF, DOCX or TXT) and uses an LLM
to turn it into a structured candidate profile used for job matching.
"""

import io
import json
import re
import logging

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger("cv_parser")


def extract_text(file_bytes, filename):
    """Return plain text content of a CV file (pdf / docx / txt)."""
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        try:
            return "\n".join(page.get_text() for page in doc)
        finally:
            doc.close()

    if name.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    if name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")

    raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")


PROFILE_PROMPT = """You are an expert technical recruiter. Read the CV below and \
return ONLY a JSON object, with no markdown formatting and no commentary, \
with exactly these keys:

{{
  "name": "the candidate's name, or empty string if not found",
  "summary": "one or two plain sentences summarising who this person is professionally",
  "experience_level": "one of: Internship, Entry, Mid, Senior",
  "target_roles": ["3 to 6 specific job titles this person should search for, most relevant first"],
  "skills": ["10 to 20 key technical and professional skills, most important first"],
  "location": "city and/or country mentioned in the CV, or empty string"
}}

CV TEXT:
{cv_text}
"""


def _strip_json_fences(text):
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    return text.strip()


def extract_profile(cv_text, client, model):
    """Use the LLM to turn raw CV text into a structured profile dict.
    Falls back to a keyword based profile if the LLM call fails."""
    try:
        response = client.chat(
            model=model,
            messages=[
                {"role": "system", "content": "You return only valid JSON, nothing else."},
                {"role": "user", "content": PROFILE_PROMPT.format(cv_text=cv_text[:8000])},
            ],
        )
        raw = _strip_json_fences(response.message.content)
        profile = json.loads(raw)
        profile.setdefault("name", "")
        profile.setdefault("summary", "")
        profile.setdefault("experience_level", "Entry")
        profile.setdefault("target_roles", [])
        profile.setdefault("skills", [])
        profile.setdefault("location", "")
        if not profile["target_roles"]:
            profile["target_roles"] = _fallback_profile(cv_text)["target_roles"]
        if not profile["skills"]:
            profile["skills"] = _fallback_profile(cv_text)["skills"]
        return profile
    except Exception as e:
        logger.warning("LLM profile extraction failed, using fallback: %s", e)
        return _fallback_profile(cv_text)


_COMMON_SKILLS = [
    "python", "java", "javascript", "typescript", "react", "node.js",
    "sql", "machine learning", "deep learning", "pytorch", "tensorflow",
    "nlp", "computer vision", "docker", "kubernetes", "aws", "gcp",
    "azure", "fastapi", "django", "flask", "streamlit", "langchain",
    "crewai", "data analysis", "pandas", "numpy", "c++", "rag", "llm",
    "html", "css", "git", "rest api", "agile",
]


def _fallback_profile(cv_text):
    """Rough keyword based fallback if the LLM call fails entirely."""
    text_l = cv_text.lower()
    found = [s for s in _COMMON_SKILLS if s in text_l]
    return {
        "name": "",
        "summary": "Profile extracted with limited information. Please review the detected roles and skills below.",
        "experience_level": "Entry",
        "target_roles": ["Software Engineer", "Machine Learning Engineer", "AI Engineer"],
        "skills": found[:15] if found else ["Software Development"],
        "location": "",
    }
